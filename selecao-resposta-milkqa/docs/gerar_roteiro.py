# -*- coding: utf-8 -*-
"""Gera o roteiro de apresentacao (docs/Roteiro_Apresentacao_AgroSele.docx)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

AZUL = RGBColor(0x1B, 0x3A, 0x5C)
VERDE = RGBColor(0x2E, 0x7D, 0x32)
CINZA = RGBColor(0x4A, 0x4A, 0x4A)

d = Document()

# estilo base
style = d.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def titulo(texto, nivel=1, cor=AZUL):
    h = d.add_heading(level=nivel)
    r = h.add_run(texto)
    r.font.color.rgb = cor
    return h

def paragrafo(texto, negrito=False, italico=False, cor=None, tamanho=11):
    p = d.add_paragraph()
    r = p.add_run(texto)
    r.bold = negrito
    r.italic = italico
    r.font.size = Pt(tamanho)
    if cor:
        r.font.color.rgb = cor
    return p

def bloco_slide(numero, titulo_slide, tempo, fala, transicao=None):
    h = d.add_heading(level=2)
    r = h.add_run(f"Slide {numero} — {titulo_slide}")
    r.font.color.rgb = VERDE

    p = d.add_paragraph()
    r = p.add_run(f"Tempo sugerido: {tempo}")
    r.italic = True
    r.font.color.rgb = CINZA
    r.font.size = Pt(10)

    p = d.add_paragraph()
    r = p.add_run("Fala: ")
    r.bold = True
    p.add_run(fala)

    if transicao:
        p = d.add_paragraph()
        r = p.add_run("Transição: ")
        r.bold = True
        r.font.color.rgb = AZUL
        p.add_run(transicao)


# ---------------------------------------------------------------------
# Capa
# ---------------------------------------------------------------------
h = d.add_heading(level=0)
r = h.add_run("Roteiro de Apresentação — AgroSele")
r.font.color.rgb = AZUL
paragrafo("Baseado em docs/Apresentacao_AgroSele.pptx (12 slides) e docs/demo_sala_aula.ipynb", italico=True, cor=CINZA)
paragrafo("Frederico Botelho Martins — PLN, Prof. Dr. Yuri Malheiros, 2026.1", cor=CINZA)

paragrafo("")
titulo("Como usar este roteiro", nivel=2, cor=AZUL)
paragrafo(
    "Tempo total estimado: aproximadamente 14-15 minutos (11 minutos de slides + "
    "3-4 minutos de demonstração ao vivo), mais tempo de perguntas. Se o slot for "
    "mais curto, os cortes sugeridos estão marcados como \"[CORTE SE FALTAR TEMPO]\" "
    "junto ao slide correspondente. A fala de cada slide é um roteiro de apoio, não "
    "um texto para ler literalmente — fale com suas palavras, os bullets do slide já "
    "estão visíveis para a audiência, seu papel é adicionar contexto, não repetir o "
    "que está escrito."
)

paragrafo("")
titulo("Preparação antes da aula", nivel=2, cor=AZUL)
p = d.add_paragraph(style="List Bullet")
p.add_run("Abra docs/demo_sala_aula.ipynb e rode a célula de Setup uma vez antes de começar "
          "(leva ~6-20s) — assim, na hora da demo, só falta rodar as células de cada método, "
          "que são instantâneas.")
p = d.add_paragraph(style="List Bullet")
p.add_run("O notebook já foi commitado com outputs reais (pergunta #42, o caso do distrator "
          "\"ração de baixo custo\"). Se a demo ao vivo falhar por qualquer motivo (sem internet, "
          "ambiente diferente), role até essas células já executadas e mostre o resultado salvo "
          "— o discurso funciona igual, só que sem o suspense de rodar na hora.")
p = d.add_paragraph(style="List Bullet")
p.add_run("Confirme que o notebook abre no mesmo notebook/laptop que será usado na apresentação "
          "(caches e checkpoints são grandes; copiar tudo na hora não é viável).")

d.add_page_break()

# ---------------------------------------------------------------------
# Slides
# ---------------------------------------------------------------------
titulo("Roteiro slide a slide", nivel=1)

bloco_slide(
    0, "Capa", "30s",
    "Cumprimente e diga em uma frase o que o sistema faz: \"um sistema que, dado uma "
    "pergunta real de um produtor rural e 50 respostas candidatas da Embrapa, aponta "
    "qual delas é a correta\". Mencione rapidamente a motivação pessoal: o tema se conecta "
    "a um software de gestão para fazendas leiteiras que você já desenvolve.",
    transicao="\"Vou contar essa história em cinco partes: por que isso importa, como testamos, "
              "o que descobrimos, uma demonstração ao vivo, e o que ainda falta.\"",
)

bloco_slide(
    1, "Motivação", "60s",
    "Conte o \"gargalo humano\": cada pergunta nova, mesmo quando semanticamente idêntica "
    "a uma já respondida por um especialista, ainda demanda trabalho manual individualizado. "
    "Isso não escala. Destaque o desafio específico do domínio: texto informal e heterogêneo, "
    "escrito por produtores sem formação técnica, com vocabulário técnico compartilhado entre "
    "respostas de assuntos diferentes — o que torna a tarefa mais difícil do que parece.",
)

bloco_slide(
    2, "Pergunta de Pesquisa", "30s",
    "Leia a pergunta de pesquisa pausadamente e deixe 2-3 segundos de silêncio depois — "
    "é densa, a audiência precisa de um instante para processar. Frise que a pergunta não "
    "é \"o modelo funciona?\", mas sim \"que estratégia de combinação compensa o custo "
    "computacional?\" — isso já prepara a audiência para o argumento central do achado principal.",
)

bloco_slide(
    3, "Dados: MilkQA", "60s",
    "Explique a mecânica da tarefa antes dos números: não é classificação em categorias "
    "fixas, é ranquear 50 respostas candidatas e escolher a mais provável. Cite os números "
    "principais (2.657 pares, splits oficiais) rapidamente, mas pare um pouco mais na última "
    "linha: perguntas longas e informais, média de 198 palavras — dê um exemplo mental rápido "
    "de como é esse texto (o exemplo do Slide 8 já mostra um).",
)

bloco_slide(
    4, "Abordagem: 6 configurações comparadas", "90s",
    "Passe pelas 6 variantes rapidamente, apenas nomeando a ideia central de cada uma — "
    "os detalhes técnicos vêm depois, no achado principal e na demo. O importante aqui é "
    "a audiência perceber a progressão de complexidade: de um método sem nenhum modelo de "
    "linguagem até fine-tuning parcial do BERTimbau.",
    transicao="[CORTE SE FALTAR TEMPO] — se precisar cortar, resuma em uma frase: \"testamos seis "
              "variações, do TF-IDF puro ao fine-tuning do BERTimbau, sempre no mesmo protocolo de "
              "avaliação\", e passe direto ao gráfico de resultados.",
)

bloco_slide(
    5, "Métricas de Avaliação", "45s",
    "Explique por que F1-macro não se aplica (é ranking, não classificação em classes fixas). "
    "Ancore visualmente com o baseline aleatório: Accuracy@1 de apenas 0,02 — isso dá à audiência "
    "uma régua mental para julgar os resultados que vêm a seguir.",
)

bloco_slide(
    6, "Resultados (gráfico de barras)", "90s",
    "Deixe a audiência olhar o gráfico por 3-4 segundos antes de falar. Depois, guie o olhar "
    "da esquerda para a direita: cada barra é uma pergunta diferente que fizemos ao problema "
    "— \"e se usarmos léxico puro?\", \"e se usarmos embeddings sem treino?\", \"e se treinarmos "
    "um classificador?\", até chegar no fine-tuning parcial.",
)

bloco_slide(
    7, "Achado Principal", "120s",
    "Este é o coração da apresentação — não tenha pressa aqui. Conte a \"virada de mesa\": "
    "TF-IDF, sem nenhum modelo de linguagem, quase dobra o BERTimbau usado cru (0,503 vs 0,277). "
    "Explique o porquê: vocabulário técnico compartilhado favorece correspondência lexical exata "
    "num domínio como esse. Em seguida, conte o episódio do fine-tuning: uma primeira tentativa, "
    "com apenas 300 das 2.307 perguntas de treino disponíveis, sugeriu que fine-tuning não "
    "compensava o custo computacional — e só repetir o experimento com o conjunto completo "
    "revelou o contrário. Enfatize a lição metodológica: uma limitação computacional pode "
    "mascarar o real potencial de uma técnica.",
)

bloco_slide(
    8, "Exemplo Concreto: Entrada e Saída", "90s",
    "Leia o caso de sucesso rapidamente (a pergunta sobre o leite azedando, resolvida "
    "corretamente). Dedique mais tempo ao caso de erro: a pergunta sobre horário de "
    "alimentação das vacas, em que o sistema escolheu por engano uma resposta sobre fórmula "
    "de ração de baixo custo. Este caso de erro é o gancho direto para a demonstração ao vivo "
    "— avise a audiência: \"vamos ver esse mesmo tipo de erro acontecer ao vivo daqui a pouco\".",
)

bloco_slide(
    9, "Demonstração ao Vivo", "15s de fala + 3-4min de demo",
    "\"Em vez de só mostrar números prontos, vamos ver isso acontecer ao vivo, com uma "
    "pergunta que vocês escolherem.\" Peça um número de 0 a 299 para a turma. Abra o "
    "notebook (Setup já deve estar rodado). Mude INDICE_PERGUNTA para o número escolhido, "
    "rode a célula do Passo 1 (mostra a pergunta), depois as células de TF-IDF, cosseno e "
    "Cross-Encoder em sequência — cada uma leva menos de 1 segundo. Comente os resultados "
    "conforme aparecem: \"olha a posição da resposta certa em cada método\", \"por que vocês "
    "acham que esse método errou aqui?\". Termine mostrando o resumo comparativo e o gabarito.",
    transicao="Plano B: se a demo travar por qualquer motivo, role até as células já executadas "
              "(pergunta #42, o mesmo padrão de erro do Slide 8) e mostre o resultado salvo — "
              "funciona igual para a discussão, só sem o suspense de rodar na hora.",
)

bloco_slide(
    10, "Limitações", "60s",
    "Passe pelos pontos com honestidade, sem se alongar — isso reforça credibilidade "
    "científica em vez de enfraquecer a apresentação. Se o tempo estiver apertado, destaque "
    "só dois: o truncamento de texto e o fato de o modelo ainda errar diante de candidatas "
    "lexicalmente densas (conecta de volta com a demo).",
)

bloco_slide(
    11, "Conclusão", "60s",
    "Feche com as três lições centrais: sinal lexical explícito e volume de dado pesam tanto "
    "quanto a arquitetura; fine-tuning parcial é o melhor resultado isolado, mas o híbrido "
    "BM25 chega perto a uma fração do custo; e o pipeline multi-estágio recupera qualidade a "
    "80% menos custo computacional — o princípio central de sistemas de busca em produção, "
    "aqui demonstrado em miniatura. Deixe o link do repositório visível por alguns segundos "
    "antes de abrir para perguntas.",
)

d.add_page_break()
titulo("Perguntas frequentes esperadas (preparação)", nivel=1)
perguntas = [
    ("Por que não usar uma LLM comercial (GPT, Gemini) direto?",
     "Custo, latência, privacidade dos dados do produtor e determinismo — decisão já "
     "justificada na Fase 1 do projeto. O sistema é pensado para rodar embarcado num "
     "software de gestão de fazenda, não como chamada a uma API externa a cada pergunta."),
    ("O TF-IDF vencer o BERT cru não invalida o uso de modelos de linguagem?",
     "Não — o BERTimbau só supera o TF-IDF depois de receber algum treinamento supervisionado "
     "(classificador de pares ou fine-tuning). O achado mostra que o embedding pré-treinado, "
     "sozinho, não garante vantagem; a vantagem vem de ajustá-lo à tarefa específica."),
    ("Por que só a última camada do BERTimbau foi descongelada no fine-tuning?",
     "Restrição de tempo de processamento em CPU sem GPU dedicada. É uma das direções de "
     "trabalho futuro explicitamente apontadas no artigo."),
    ("O Cross-Encoder usa \"atenção cruzada\" como no Transformer original (encoder-decoder)?",
     "Não exatamente — tecnicamente é self-attention operando conjuntamente sobre a sequência "
     "concatenada pergunta+candidata (query, key e value vêm todos da mesma sequência), "
     "diferente da cross-attention entre encoder e decoder de arquiteturas sequence-to-sequence. "
     "O nome \"Cross-Encoder\" vem da literatura (Nogueira & Cho, 2019) e se refere a processar "
     "o par conjuntamente, não ao mecanismo de atenção em si."),
]
for pergunta, resposta in perguntas:
    p = d.add_paragraph()
    r = p.add_run("P: ")
    r.bold = True
    r.font.color.rgb = VERDE
    p.add_run(pergunta)
    p2 = d.add_paragraph()
    r2 = p2.add_run("R: ")
    r2.bold = True
    r2.font.color.rgb = AZUL
    p2.add_run(resposta)
    d.add_paragraph("")

d.save("Roteiro_Apresentacao_AgroSele.docx")
print("Roteiro gerado: Roteiro_Apresentacao_AgroSele.docx")
